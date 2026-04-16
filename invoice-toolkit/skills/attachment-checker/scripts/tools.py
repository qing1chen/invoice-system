"""
附件检查工具函数集（供 LLM 编排调用）

本模块提供 SKILL.md 中定义的所有工具函数的 Python 实现。
每个函数都是无状态的纯工具——接收参数，执行文件IO/数据库/OCR操作，返回结果。
所有决策逻辑由 LLM 在 SKILL.md 指导下完成，本模块不做任何业务判断。

使用方式：
    - 作为 MCP Server 的 tool handler 注册
    - 或由 LLM 通过 bash 脚本调用 `python -m scripts.tools <tool_name> <args_json>`
"""

from __future__ import annotations

import base64
import copy
import io
import json
import logging
import math
import os
import random
import re
import shutil
import subprocess
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

import requests

logger = logging.getLogger(__name__)

# ── 全局状态（在单次运行中保持）──
_used_source_attachments: Set[str] = set()

# ── 需要删除的提示文本 ─────────────────────────────────────
_HINT_PATTERN = re.compile(
    r'[（(]\s*提示[：:].+?就是这里要填的人数\s*[）)]',
    re.DOTALL,
)

# ── 报销系统可直接上传的格式（无需转换）──
_UPLOAD_OK_EXTS = {
    ".pdf", ".jpg", ".jpeg", ".png", ".gif", ".bmp",
    ".tiff", ".tif", ".webp",
}

_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".gif", ".webp"}

# ── LibreOffice 是否可用（None=未检测, True/False=已检测）──
_soffice_available: Optional[bool] = None


# =========================================================================
# 配置与初始化
# =========================================================================

def get_config(settings) -> Dict[str, Any]:
    """获取系统配置信息。"""
    from invoice_toolkit.database import get_invoice_db
    paths = settings.paths

    # 从 rules.md 解析可用类别
    categories = list(_parse_categories_from_rules(paths).keys()) if hasattr(paths, 'attachment_skill_dir') else []
    if not categories:
        categories = ["打车", "出差", "加班餐", "打印", "快递", "材料"]

    return {
        "name_list": settings.NAME_LIST,
        "source_root": str(paths.source_root),
        "invoice_root": str(paths.invoice_root),
        "categories": categories,
        "overtime_meal_output_dir": str(paths.overtime_meal_output_dir),
        "cache_dir": str(paths.cache_dir),
    }


def _parse_categories_from_rules(paths) -> Dict[str, Dict]:
    """从 rules.md 解析类别列表（简化版，完整版在 mcp_server 中）。"""
    rules_path = Path(paths.attachment_skill_dir) / "references" / "rules.md" if hasattr(paths, 'attachment_skill_dir') else None
    if not rules_path or not rules_path.exists():
        return {}
    text = rules_path.read_text(encoding="utf-8")
    categories = {}
    for match in re.finditer(r'^## (\S+)', text, re.MULTILINE):
        name = match.group(1)
        if "通用" not in name and "匹配策略" not in name and "异常" not in name and "文件名" not in name:
            categories[name] = {}
    return categories


# =========================================================================
# 文件收集
# =========================================================================

def get_ocr_names(settings) -> Dict[str, Any]:
    """获取已被 OCR 识别为发票的文件名集合。"""
    from invoice_toolkit.database import get_invoice_db
    invoice_db = get_invoice_db(settings)

    try:
        db_df = invoice_db.to_dataframe()
        if not db_df.empty and "旧文件名" in db_df.columns:
            names = list(db_df["旧文件名"].dropna())
            return {"ocr_names": names}
    except Exception:
        pass

    # 回退：从 Excel
    import pandas as pd
    p = settings.paths.ocr_excel
    if p.exists():
        try:
            df = pd.read_excel(str(p))
            names = list(df["旧文件名"].dropna()) if "旧文件名" in df.columns else []
            return {"ocr_names": names}
        except Exception:
            pass

    return {"ocr_names": []}


def collect_files(settings, category: str) -> List[Dict[str, Any]]:
    """收集指定类别目录下的所有文件。"""
    from invoice_toolkit.database import get_invoice_db
    invoice_db = get_invoice_db(settings)
    paths = settings.paths

    def _person(filename):
        for n in settings.NAME_LIST:
            if n in filename:
                return n
        return ""

    # 优先从数据库读取
    try:
        cls_df = invoice_db.get_classification()
        if not cls_df.empty:
            cat = cls_df[cls_df["category"] == category]
            if not cat.empty:
                files = []
                for _, row in cat.iterrows():
                    par = str(row.get("parent", ""))
                    per = par.split(os.sep)[0] if par and par != "." else ""
                    files.append({
                        "name": row["name"],
                        "full_path": str(row.get("full_path", "")),
                        "parent": par,
                        "person": per,
                    })
                if files:
                    # 补充同人名的非发票文件
                    ocr_data = get_ocr_names(settings)
                    ocr_set = set(ocr_data["ocr_names"])
                    existing = {f["name"] for f in files}
                    persons = {f["person"] for f in files if f["person"]}
                    for _, row in cls_df.iterrows():
                        name = row["name"]
                        if name in existing:
                            continue
                        par = str(row.get("parent", ""))
                        per = par.split(os.sep)[0] if par and par != "." else ""
                        if per in persons and name not in ocr_set:
                            files.append({
                                "name": name,
                                "full_path": str(row.get("full_path", "")),
                                "parent": par,
                                "person": per,
                            })
                    return files
    except Exception:
        pass

    # 回退：文件夹扫描
    cd = paths.invoice_root / category
    if cd.exists():
        files = []
        for f in cd.rglob("*"):
            if f.is_file():
                files.append({
                    "name": f.name,
                    "full_path": str(f),
                    "parent": str(f.parent.relative_to(cd)),
                    "person": _person(f.name),
                })
        return files

    return []


def collect_source_candidates(settings, person: str) -> List[Dict[str, Any]]:
    """从来源目录收集候选附件。"""
    global _used_source_attachments

    if not person:
        return []

    person_dir = settings.paths.source_root / person
    if not person_dir.exists():
        return []

    ocr_data = get_ocr_names(settings)
    ocr_set = set(ocr_data["ocr_names"])

    candidates = []
    for f in person_dir.rglob("*"):
        if not f.is_file():
            continue
        if f.name in ocr_set:
            continue
        if str(f) in _used_source_attachments:
            continue
        candidates.append({
            "name": f.name,
            "full_path": str(f),
            "parent": str(f.parent.relative_to(settings.paths.source_root)),
            "person": person,
        })

    return candidates


# =========================================================================
# 发票详情查询
# =========================================================================

def lookup_invoice_details(settings, filename: str) -> Dict[str, Any]:
    """从数据库查询发票 OCR 详情。"""
    from invoice_toolkit.database import get_invoice_db
    invoice_db = get_invoice_db(settings)
    return invoice_db.lookup_invoice_details(filename)


# =========================================================================
# 附件 OCR 文字提取
# =========================================================================

_GENERAL_OCR_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/general_basic"
_ocr_token_cache: Optional[str] = None


def extract_attachment_text(settings, filepath: str) -> Optional[Dict[str, Any]]:
    """统一的附件文字提取（截图+OCR）。"""
    path = Path(filepath)
    if not path.exists():
        return None

    suffix = path.suffix.lower()

    try:
        if suffix in _IMAGE_EXTENSIONS:
            text = _ocr_image(settings, path)
            method = "image_ocr"
        elif suffix == ".pdf":
            text = _ocr_pdf(settings, path)
            method = "pdf_ocr"
        elif suffix in (".doc", ".docx"):
            text = _read_doc_text(path)
            method = "docx_text"
            if not text:
                pdf_path = _convert_to_pdf_soffice(settings, path)
                if pdf_path:
                    text = _ocr_pdf(settings, pdf_path)
                    method = "doc_via_pdf_ocr"
        else:
            text = _ocr_image(settings, path)
            method = "fallback_ocr"

        if not text:
            return None

        truncated = len(text) > 3000
        return {
            "text": text[:3000],
            "truncated": truncated,
            "method": method,
        }

    except Exception as exc:
        logger.warning("OCR 提取失败: %s — %s", path.name, exc)
        return None


def _ocr_image(settings, path: Path) -> Optional[str]:
    """对图片文件调用百度 OCR 通用文字识别。"""
    try:
        with open(path, "rb") as f:
            img_bytes = f.read()
        return _ocr_bytes(settings, img_bytes)
    except Exception:
        return None


def _ocr_pdf(settings, pdf_path: Path) -> Optional[str]:
    """将 PDF 渲染为图片后 OCR。"""
    try:
        import fitz
    except ImportError:
        return None

    try:
        doc = fitz.open(str(pdf_path))
        parts = []
        for i, page in enumerate(doc):
            if i >= 3:
                parts.append(f"（共 {len(doc)} 页，仅识别前 3 页）")
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            text = _ocr_bytes(settings, pix.tobytes("png"))
            if text:
                parts.append(text)
        doc.close()
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def _ocr_bytes(settings, img_bytes: bytes) -> Optional[str]:
    """调用百度 OCR 通用文字识别 API。"""
    global _ocr_token_cache

    if not _ocr_token_cache:
        try:
            ocr_settings = settings.ocr
            if not ocr_settings.api_key or not ocr_settings.secret_key:
                return None
            resp = requests.get(
                ocr_settings.token_url,
                params={
                    "grant_type": "client_credentials",
                    "client_id": ocr_settings.api_key,
                    "client_secret": ocr_settings.secret_key,
                },
                timeout=30,
            )
            data = resp.json()
            _ocr_token_cache = data.get("access_token")
            if not _ocr_token_cache:
                return None
        except Exception:
            return None

    try:
        image_b64 = base64.b64encode(img_bytes).decode("utf-8")
        resp = requests.post(
            f"{_GENERAL_OCR_URL}?access_token={_ocr_token_cache}",
            data={"image": image_b64},
            headers={"User-Agent": "invoice-toolkit"},
            timeout=30,
        )
        data = resp.json()
        words_result = data.get("words_result", [])
        if not words_result:
            return None
        return "\n".join(item.get("words", "") for item in words_result)
    except Exception:
        return None


def _read_doc_text(path: Path) -> Optional[str]:
    """用 python-docx 直接读取 doc/docx 文本。"""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return None
    try:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)
        return "\n".join(parts) if parts else None
    except Exception:
        return None


# =========================================================================
# LibreOffice 检测与文档转换
# =========================================================================

def _check_soffice() -> bool:
    """检测 LibreOffice 是否可用（只检测一次）。"""
    global _soffice_available
    if _soffice_available is not None:
        return _soffice_available
    for exe in ("soffice", "libreoffice"):
        try:
            subprocess.run(
                [exe, "--version"], capture_output=True, timeout=10,
            )
            _soffice_available = True
            logger.info("检测到 LibreOffice (%s)，附件转 PDF 将使用 soffice", exe)
            return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    _soffice_available = False
    logger.info("未检测到 LibreOffice，附件转 PDF 将使用 PyMuPDF 纯 Python 方案")
    return False


def _run_soffice_convert_pdf(src_path: Path, outdir: Path) -> Optional[Path]:
    """调用 LibreOffice 将文档转换为 PDF。"""
    expected = outdir / f"{src_path.stem}.pdf"

    # 清理 LibreOffice 锁文件
    lock_dir = Path.home() / ".config" / "libreoffice"
    lock_file = lock_dir / ".~lock.localhost#"
    if lock_file.exists():
        try:
            lock_file.unlink()
        except OSError:
            pass

    for exe in ("soffice", "libreoffice"):
        try:
            subprocess.run(
                [exe, "--headless", "--convert-to", "pdf",
                 "--outdir", str(outdir), str(src_path)],
                check=True, capture_output=True, timeout=60,
                env={**os.environ, "HOME": str(Path.home())},
            )
            if expected.exists():
                return expected
            # 查找替代输出文件名
            for f in outdir.glob("*.pdf"):
                if src_path.stem.lower() in f.stem.lower():
                    return f
        except FileNotFoundError:
            continue
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("  soffice PDF 转换异常: %s — %s", src_path.name, exc)
            return None

    return None


def _convert_to_pdf_soffice(settings, src_path: Path) -> Optional[Path]:
    """将 doc/docx 转换为 PDF（缓存到 cache/attachment_pdf/）。

    仅使用 LibreOffice，不含 PyMuPDF 回退。用于 OCR 前的格式转换。
    """
    cache_dir = settings.paths.cache_dir / "attachment_pdf"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out = cache_dir / f"{src_path.stem}.pdf"
    if out.exists():
        return out
    for exe in ("soffice", "libreoffice"):
        try:
            subprocess.run(
                [exe, "--headless", "--convert-to", "pdf",
                 "--outdir", str(cache_dir), str(src_path)],
                check=True, capture_output=True, timeout=60,
            )
            if out.exists():
                return out
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue
    return None


def _docx_to_pdf_pymupdf(src_path: Path, out_path: Path) -> Optional[Path]:
    """纯 Python 方案：python-docx 读取 + PyMuPDF 写 PDF。

    不依赖 LibreOffice，利用已有依赖 PyMuPDF（内置 CJK 字体）。
    适用于结构简单的加班餐情况说明等文档。
    """
    suffix = src_path.suffix.lower()
    if suffix not in (".docx", ".doc", ".txt"):
        logger.info("  PyMuPDF 回退不支持 %s 格式，跳过: %s", suffix, src_path.name)
        return None

    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("  PyMuPDF 未安装，无法进行纯 Python PDF 转换")
        return None

    # ── 提取文本内容 ──
    lines: List[str] = []
    if suffix == ".txt":
        try:
            lines = src_path.read_text(encoding="utf-8").splitlines()
        except Exception:
            lines = src_path.read_text(encoding="gbk", errors="replace").splitlines()
    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(src_path))
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    line = " | ".join(cells)
                    if line.replace("|", "").strip():
                        lines.append(line)
        except Exception as exc:
            logger.warning("  python-docx 读取失败: %s — %s", src_path.name, exc)
            return None

    if not lines:
        logger.warning("  文档内容为空，跳过 PDF 转换: %s", src_path.name)
        return None

    # ── 用 PyMuPDF 生成 PDF（内置 CJK 字体，无需额外安装）──
    try:
        pdf_doc = fitz.open()
        # A4 尺寸
        width, height = 595.28, 841.89
        margin = 56.7  # 约 2cm 边距
        font_size = 11
        line_height = font_size * 1.8
        # PyMuPDF 内置中文字体
        fontname = "china-s"

        y = margin
        page = pdf_doc.new_page(width=width, height=height)

        for line_text in lines:
            # 自动换行：按可用宽度估算每行字符数
            max_chars = int((width - 2 * margin) / (font_size * 0.6))
            sub_lines = []
            while len(line_text) > max_chars:
                sub_lines.append(line_text[:max_chars])
                line_text = line_text[max_chars:]
            sub_lines.append(line_text)

            for sub in sub_lines:
                if y + line_height > height - margin:
                    page = pdf_doc.new_page(width=width, height=height)
                    y = margin
                page.insert_text(
                    fitz.Point(margin, y),
                    sub,
                    fontname=fontname,
                    fontsize=font_size,
                )
                y += line_height

        pdf_doc.save(str(out_path))
        pdf_doc.close()
        logger.info("  ✓ PyMuPDF 纯 Python 转换成功: %s", out_path.name)
        return out_path

    except Exception as exc:
        logger.error("  PyMuPDF PDF 生成失败: %s — %s", src_path.name, exc)
        return None


def _convert_to_pdf_cached(settings, src_path: Path) -> Optional[Path]:
    """将任意文档转为 PDF，缓存到 cache/attachment_pdf/。

    策略（与旧版 AttachmentChecker._convert_to_pdf_cached 一致）：
      1. 缓存命中 → 直接返回
      2. 失败标记存在 → 跳过（避免重复尝试）
      3. LibreOffice 可用 → soffice --convert-to pdf
      4. 回退 → PyMuPDF + python-docx 纯 Python 转换
      5. 全部失败 → 写入失败标记
    """
    cache_dir = settings.paths.cache_dir / "attachment_pdf"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out = cache_dir / f"{src_path.stem}.pdf"

    logger.info("  [PDF转换] 开始: %s → %s", src_path.name, out.name)

    # 1. 缓存命中
    if out.exists():
        logger.info("  [PDF转换] 缓存命中: %s", out.name)
        return out

    # 2. 失败标记（避免同一文件反复尝试）
    fail_marker = cache_dir / f"{src_path.stem}.pdf.failed"
    if fail_marker.exists():
        logger.info("  [PDF转换] 存在失败标记，跳过: %s", src_path.name)
        return None

    # 3. 尝试 LibreOffice
    pdf = None
    if _check_soffice():
        pdf = _run_soffice_convert_pdf(src_path, cache_dir)
        if pdf:
            logger.info("  [PDF转换] LibreOffice 成功: %s", pdf.name)
        else:
            logger.info("  [PDF转换] LibreOffice 失败，尝试 PyMuPDF 回退")

    # 4. 回退 PyMuPDF 纯 Python
    if not pdf:
        pdf = _docx_to_pdf_pymupdf(src_path, out)

    if pdf and pdf.exists():
        logger.info("  ✓ 附件已转为 PDF: %s → %s", src_path.name, pdf.name)
        # 清除可能遗留的失败标记
        fail_marker.unlink(missing_ok=True)
        return pdf

    # 全部失败 → 写入失败标记
    fail_marker.write_text(f"failed at {datetime.now().isoformat()}")
    logger.warning("  附件转 PDF 失败: %s（已标记，不再重试）", src_path.name)
    return None


def _ensure_upload_format(settings, att_path: str) -> str:
    """确保附件路径指向 PDF 或图片，否则转换为 PDF。

    输入：逗号分隔的路径字符串（可能多个附件）
    输出：转换后的逗号分隔路径字符串

    与旧版 AttachmentChecker._ensure_upload_format 逻辑一致。
    """
    if not att_path:
        return att_path
    parts = [p.strip() for p in att_path.split(",") if p.strip()]
    converted = []
    seen: Set[str] = set()  # 去重：同一路径只转换一次
    for p in parts:
        if p in seen:
            cache_dir = settings.paths.cache_dir / "attachment_pdf"
            cached = cache_dir / f"{Path(p).stem}.pdf"
            converted.append(str(cached) if cached.exists() else p)
            continue
        seen.add(p)
        path = Path(p)
        if path.suffix.lower() in _UPLOAD_OK_EXTS:
            converted.append(p)
            continue

        # 文件不存在时：尝试按文件名在 invoice_root 下搜索
        if not path.exists():
            found_path = None
            for f in settings.paths.invoice_root.rglob(path.name):
                if f.is_file():
                    found_path = f
                    break
            if found_path:
                logger.info(
                    "  [格式转换] 路径不存在，通过搜索找到: %s → %s",
                    p, found_path,
                )
                path = found_path
            else:
                logger.info(
                    "  [格式转换] 文件不存在，跳过转换: %s", p,
                )
                converted.append(p)
                continue

        logger.info("  [格式转换] 开始转换: %s (exists=%s)", path, path.exists())
        pdf = _convert_to_pdf_cached(settings, path)
        if pdf:
            converted.append(str(pdf))
        else:
            converted.append(str(path))
    return ",".join(converted)


# =========================================================================
# 文件操作
# =========================================================================

def copy_file(src: str, dst_dir: str, mark_used: bool = True) -> Dict[str, Any]:
    """复制文件到目标目录。"""
    global _used_source_attachments

    src_path = Path(src)
    dst_dir_path = Path(dst_dir)
    dst_dir_path.mkdir(parents=True, exist_ok=True)
    dst = dst_dir_path / src_path.name

    # 处理重名
    if dst.exists():
        stem, ext = os.path.splitext(src_path.name)
        i = 1
        while dst.exists():
            dst = dst_dir_path / f"{stem}_{i}{ext}"
            i += 1

    try:
        shutil.copy2(str(src_path), str(dst))
        if mark_used:
            _used_source_attachments.add(src)
        return {"success": True, "dst_path": str(dst), "dst_name": dst.name}
    except Exception as exc:
        return {"success": False, "error": str(exc)}


def backup_file(settings, filepath: str, delete_original: bool = True) -> Dict[str, Any]:
    """备份文件到 cache/attachment_backup/。"""
    path = Path(filepath)
    if not path.exists():
        return {"success": False, "error": "文件不存在"}

    backup_dir = settings.paths.cache_dir / "attachment_backup"
    backup_dir.mkdir(parents=True, exist_ok=True)
    dst = backup_dir / path.name

    if dst.exists():
        stem, suffix = dst.stem, dst.suffix
        i = 1
        while dst.exists():
            dst = backup_dir / f"{stem}_{i}{suffix}"
            i += 1

    try:
        shutil.copy2(str(path), str(dst))
        if delete_original:
            path.unlink()
        return {"success": True, "backup_path": str(dst)}
    except Exception as exc:
        return {"success": False, "error": str(exc)}


# =========================================================================
# 加班餐文档生成/修复/合并
# =========================================================================

def _set_run_text(para, new_text: str):
    """设置段落文本为 new_text，同时保留第一个 run 的字体格式。"""
    if not para.runs:
        para.text = new_text
        return

    first_run = para.runs[0]
    orig_font_name = first_run.font.name
    orig_font_size = first_run.font.size
    orig_bold = first_run.bold
    orig_italic = first_run.italic

    first_run.text = new_text
    for run in para.runs[1:]:
        run.text = ""

    if orig_font_name:
        first_run.font.name = orig_font_name
    if orig_font_size:
        first_run.font.size = orig_font_size
    if orig_bold is not None:
        first_run.bold = orig_bold
    if orig_italic is not None:
        first_run.italic = orig_italic


def generate_meal_doc(settings, person: str, amount: float, seller: str,
                      commodity: str, invoice_filename: str,
                      name_list: List[str],
                      reason: Optional[str] = None) -> Dict[str, Any]:
    """生成加班餐情况说明 docx。"""
    if not reason:
        reason = "加班整理资料申报科研项目，订购外卖"

    required = max(math.ceil(amount / 30), 1) if amount > 0 else 1
    used = set()
    if person and person in name_list:
        used.add(person)

    avail = [n for n in name_list if n not in used]
    random.shuffle(avail)
    while len(used) < required and avail:
        used.add(avail.pop(0))

    # 排序：报销人排首位
    ordered = sorted(used, key=lambda n: name_list.index(n) if n in name_list else 999)
    if person and person in ordered:
        ordered.remove(person)
        ordered.insert(0, person)

    persons_text = f"{len(ordered)}人，{'，'.join(ordered)}"

    output_dir = Path(settings.paths.overtime_meal_output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    amount_str = str(amount) if amount else "0"
    fname = f"{person}+{amount_str}+加班餐.docx"
    if len(fname) > 200:
        fname = f"{person}+{amount_str}+加班餐.docx"
    out_path = output_dir / fname

    if out_path.exists():
        try:
            out_path.unlink()
        except OSError:
            pass

    # 从零创建文档
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return {"success": False, "error": "python-docx 未安装"}

    try:
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(11)

        doc.add_paragraph("附件1")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("山东大学科研业务专项经费使用说明表")
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "宋体"
        doc.add_paragraph("单位公章:")

        event_text = f"{reason}，{persons_text}"
        table = doc.add_table(rows=3, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "序号"
        table.rows[0].cells[1].text = "事由"
        table.rows[0].cells[2].text = "接待费/燃油费"
        table.rows[0].cells[3].text = "金额（元）"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = event_text
        table.rows[1].cells[2].text = "接待费"
        table.rows[1].cells[3].text = amount_str
        table.rows[2].cells[0].text = "合计"
        table.rows[2].cells[3].text = amount_str

        doc.add_paragraph("")

        p1 = doc.add_paragraph()
        run1 = p1.add_run("兹证明该事项真实有效，本人愿意为此承担责任。")
        run1.font.name = "宋体"
        run1.font.size = Pt(11)

        doc.add_paragraph("")

        p2 = doc.add_paragraph()
        run2 = p2.add_run("项目负责人（签名）：")
        run2.font.name = "宋体"
        run2.font.size = Pt(11)

        doc.add_paragraph("")

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = p3.add_run(datetime.now().strftime('%Y年%m月%d日'))
        run3.font.name = "宋体"
        run3.font.size = Pt(11)

        doc.add_paragraph("")

        p4 = doc.add_paragraph()
        run4 = p4.add_run(
            "项目负责人是科研经费使用的直接责任人，"
            "对经费使用的真实性、合理性及有效性承担经济与法律责任。"
        )
        run4.font.name = "宋体"
        run4.font.size = Pt(11)

        p5 = doc.add_paragraph()
        run5 = p5.add_run("——《山东大学科研经费使用与报销管理规定》")
        run5.font.name = "宋体"
        run5.font.size = Pt(11)

        doc.save(str(out_path))
        return {
            "success": True,
            "generated_path": str(out_path),
            "persons_text": persons_text,
        }
    except Exception as exc:
        return {"success": False, "error": str(exc)}


def fix_meal_doc(settings, original_path: str, invoice_filename: str,
                 person: str, amount: float, target_persons: List[str],
                 required_count: int, reason_text: Optional[str] = None) -> Dict[str, Any]:
    """修复加班餐情况说明（实质是备份原文件后重新生成）。"""
    # 备份原文件
    backup_result = backup_file(settings, original_path, delete_original=True)
    if not backup_result["success"]:
        return {"success": False, "error": f"备份失败: {backup_result['error']}"}

    # 重新生成
    name_list = settings.NAME_LIST
    result = generate_meal_doc(
        settings, person, amount,
        seller="", commodity="餐饮",
        invoice_filename=invoice_filename,
        name_list=name_list,
        reason=reason_text,
    )

    if result["success"]:
        result["fix_msg"] = f"已重新生成，人名: {', '.join(target_persons[:5])}"

    return result


# =========================================================================
# 加班餐文件合并（从旧版 AttachmentChecker 完整移植）
# =========================================================================

def _extract_person_from_filename(filename: str, name_list: List[str]) -> str:
    """从文件名提取人名（使用 NAME_LIST 匹配）。"""
    for n in name_list:
        if n in filename:
            return n
    parts = filename.split("+")
    if parts:
        return parts[0]
    return ""


def _extract_amount_from_filename(filename: str) -> float:
    """从文件名中提取金额（第一个 '+' 分隔的数字部分）。"""
    parts = filename.replace(".docx", "").replace(".doc", "").split("+")
    for part in parts:
        try:
            v = float(part)
            if v > 0:
                return v
        except ValueError:
            continue
    return 0.0


def _build_merge_groups(person_files: Dict[str, List[Path]],
                        max_per_group: int = 6) -> List[List[Path]]:
    """
    将按人名分组的文件列表组合成合并分组。

    算法：First Fit Decreasing 装箱（与旧版一致）
      - 将每个人的文件视为不可分割的整体
      - 按文件数从多到少排序，依次尝试放入已有分组
      - 放不下则新开一组
      - 若某人的文件数超过 max_per_group，则单独拆分
    """
    batches = [
        (person, files)
        for person, files in sorted(person_files.items())
    ]
    batches.sort(key=lambda x: -len(x[1]))

    groups: List[List[Path]] = []

    for person, files in batches:
        # 极端情况：一个人的文件超过上限，必须拆分
        if len(files) > max_per_group:
            for i in range(0, len(files), max_per_group):
                groups.append(list(files[i:i + max_per_group]))
            continue

        # 尝试将该人的全部文件放入已有分组（First Fit）
        placed = False
        for group in groups:
            if len(group) + len(files) <= max_per_group:
                group.extend(files)
                placed = True
                break

        if not placed:
            groups.append(list(files))

    return groups


def _build_merged_filename(group: List[Path], name_list: List[str]) -> str:
    """
    根据合并分组中的文件构建合并后的文件名。

    命名格式：人名1、人名2、人名3+总金额+加班餐报销说明.docx
    """
    # 提取所有人名（去重，保持 NAME_LIST 顺序）
    persons_seen: Set[str] = set()
    persons_ordered: List[str] = []
    for f in group:
        person = _extract_person_from_filename(f.name, name_list)
        if person and person not in persons_seen:
            persons_seen.add(person)
            persons_ordered.append(person)

    # 按 NAME_LIST 原始顺序排序
    persons_ordered.sort(
        key=lambda n: name_list.index(n) if n in name_list else 999
    )

    # 计算总金额
    total_amount = 0.0
    for f in group:
        amt = _extract_amount_from_filename(f.name)
        total_amount += amt
    total_amount_str = str(round(total_amount, 2))
    if total_amount_str.endswith(".0"):
        total_amount_str = total_amount_str[:-2]

    # 构建文件名
    persons_str = "、".join(persons_ordered) if persons_ordered else "未知"
    fname = f"{persons_str}+{total_amount_str}+加班餐报销说明.docx"

    # 文件名过长时逐步缩减人名
    if len(fname) > 200:
        if len(persons_ordered) > 3:
            persons_str = "、".join(persons_ordered[:3]) + "等"
            fname = f"{persons_str}+{total_amount_str}+加班餐报销说明.docx"
    if len(fname) > 200:
        fname = f"{persons_ordered[0]}等+{total_amount_str}+加班餐报销说明.docx"

    return fname


def _extract_data_rows_from_table(table) -> List[Dict]:
    """从加班餐说明表格中提取纯数据行（表头和合计之间的行）。

    与旧版 AttachmentChecker._extract_data_rows_from_table 一致。
    """
    total_row_idx = None
    for i, row in enumerate(table.rows):
        cell0 = row.cells[0].text.strip()
        if "合计" in cell0:
            total_row_idx = i
            break

    if total_row_idx is None:
        return []

    data_rows = []
    for i in range(1, total_row_idx):
        row = table.rows[i]
        cell0 = row.cells[0].text.strip()

        if not cell0 or not re.match(r'^\d+$', cell0):
            continue

        data_rows.append({
            "reason": row.cells[1].text.strip(),
            "fee_type": row.cells[2].text.strip(),
            "amount": row.cells[3].text.strip(),
        })

    return data_rows


def _merge_meal_docx_files(files: List[Path], output: Path) -> bool:
    """将多个加班餐说明 .docx 合并为一个文件。

    与旧版 AttachmentChecker._merge_meal_docx_files 完全一致：
    - 使用低级 OxmlElement 操作插入行，保留表头格式
    - 字体大小设为 24 半磅（12pt）
    - 在合计行上方插入新数据行
    - 重新计算合计金额
    - 清理提示文本
    """
    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.error("需要安装 python-docx")
        return False

    if not files:
        return False

    try:
        all_data_rows = []

        for fpath in files:
            try:
                doc = DocxDocument(str(fpath))
                if not doc.tables:
                    continue
                rows = _extract_data_rows_from_table(doc.tables[0])
                all_data_rows.extend(rows)
            except Exception as exc:
                logger.warning("  读取文件失败，跳过: %s — %s",
                               fpath.name, exc)

        if not all_data_rows:
            logger.warning("没有提取到任何数据行")
            return False

        # 重排序号
        for idx, row_data in enumerate(all_data_rows, 1):
            row_data["seq"] = str(idx)

        # 以第一个文件为基础文档
        base_doc = DocxDocument(str(files[0]))
        if not base_doc.tables:
            return False

        base_table = base_doc.tables[0]
        tbl_element = base_table._tbl

        # 找到合计行
        total_row_idx = None
        for i, row in enumerate(base_table.rows):
            if "合计" in row.cells[0].text.strip():
                total_row_idx = i
                break
        if total_row_idx is None:
            return False

        # 删除基础文档中的旧数据行（保留表头和合计行）
        rows_to_remove = []
        for i in range(1, total_row_idx):
            rows_to_remove.append(base_table.rows[i]._tr)
        for tr in rows_to_remove:
            tbl_element.remove(tr)

        # 重新定位合计行（删除后索引已变）
        total_tr = None
        for row in base_table.rows:
            if "合计" in row.cells[0].text.strip():
                total_tr = row._tr
                break
        if total_tr is None:
            return False

        header_tr = base_table.rows[0]._tr

        # 在合计行上方插入新数据行
        total_amount = 0.0
        for data in all_data_rows:
            new_tr = OxmlElement('w:tr')

            # 复制表头行的行属性
            tr_pr = header_tr.find(qn('w:trPr'))
            if tr_pr is not None:
                new_tr.append(copy.deepcopy(tr_pr))

            cell_texts = [data["seq"], data["reason"],
                          data["fee_type"], data["amount"]]

            for ci, cell_text in enumerate(cell_texts):
                new_tc = OxmlElement('w:tc')

                # 复制表头对应列的单元格属性（宽度等）
                header_tc = header_tr.findall(qn('w:tc'))[ci]
                tc_pr = header_tc.find(qn('w:tcPr'))
                if tc_pr is not None:
                    new_tc.append(copy.deepcopy(tc_pr))

                new_p = OxmlElement('w:p')
                new_r = OxmlElement('w:r')

                # 从表头提取字体格式
                new_rPr = None
                header_p = header_tc.findall(qn('w:p'))
                if header_p:
                    header_runs = header_p[0].findall(qn('w:r'))
                    if header_runs:
                        rPr = header_runs[0].find(qn('w:rPr'))
                        if rPr is not None:
                            new_rPr = copy.deepcopy(rPr)

                if new_rPr is None:
                    new_rPr = OxmlElement('w:rPr')

                # 强制设置字体大小为 24 半磅（12pt）
                for old_sz in new_rPr.findall(qn('w:sz')):
                    new_rPr.remove(old_sz)
                for old_sz in new_rPr.findall(qn('w:szCs')):
                    new_rPr.remove(old_sz)
                sz_elem = OxmlElement('w:sz')
                sz_elem.set(qn('w:val'), '24')
                szCs_elem = OxmlElement('w:szCs')
                szCs_elem.set(qn('w:val'), '24')
                new_rPr.append(sz_elem)
                new_rPr.append(szCs_elem)

                new_r.append(new_rPr)

                new_t = OxmlElement('w:t')
                new_t.set(qn('xml:space'), 'preserve')
                new_t.text = cell_text
                new_r.append(new_t)
                new_p.append(new_r)
                new_tc.append(new_p)
                new_tr.append(new_tc)

            # 在合计行之前插入
            tbl_element.insert(
                list(tbl_element).index(total_tr),
                new_tr,
            )

            try:
                total_amount += float(data["amount"])
            except (ValueError, TypeError):
                pass

        # 更新合计行金额
        for row in base_table.rows:
            if "合计" in row.cells[0].text.strip():
                amt_cell = row.cells[3]
                for para in amt_cell.paragraphs:
                    if para.runs:
                        _set_run_text(para, str(round(total_amount, 2)))
                    else:
                        para.text = str(round(total_amount, 2))
                break

        # 清理提示文本
        for table in base_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full = "".join(run.text for run in para.runs)
                        if _HINT_PATTERN.search(full):
                            cleaned = _HINT_PATTERN.sub('', full).strip()
                            _set_run_text(para, cleaned)

        base_doc.save(str(output))
        return True

    except Exception as exc:
        logger.error("合并文档失败: %s", exc)
        return False


def merge_meal_docs(settings, generated_files: List[str]) -> Dict[str, Any]:
    """合并多个加班餐说明文件。

    与旧版 AttachmentChecker.merge_meal_docs 逻辑一致：
    1. 只合并显式传入的本次生成文件（不 glob 全目录）
    2. 按人名分组，使用 First Fit Decreasing 装箱算法
    3. 合并方式：在表格「合计」行上方插入新数据行（低级 OxmlElement 操作）
    4. 合并后备份并删除已合并的散件
    5. 合并后文件命名：人名1、人名2、人名3+总金额+加班餐报销说明.docx
    """
    name_list = settings.NAME_LIST
    output_dir = Path(settings.paths.overtime_meal_output_dir)

    paths = [Path(f) for f in generated_files if Path(f).exists()]
    if len(paths) <= 1:
        return {
            "merged_paths": [str(p) for p in paths],
            "merge_map": {},
            "success": True,
        }

    # 按人名分组
    person_files: Dict[str, List[Path]] = defaultdict(list)
    for f in sorted(paths):
        person = _extract_person_from_filename(f.name, name_list)
        person_files[person].append(f)

    # 使用 First Fit Decreasing 装箱分组
    groups = _build_merge_groups(person_files, max_per_group=6)

    logger.info("加班餐合并: %d 个生成文件 → %d 组", len(paths), len(groups))

    backup_dir = settings.paths.cache_dir / "meal_merge_backup"
    backup_dir.mkdir(parents=True, exist_ok=True)

    merged_paths: List[str] = []
    merge_map: Dict[str, str] = {}
    files_to_delete: List[Path] = []

    for idx, group in enumerate(groups, 1):
        if len(group) <= 1:
            merged_paths.append(str(group[0]))
            merge_map[str(group[0])] = str(group[0])
            merge_map[group[0].name] = str(group[0])
            continue

        # 构建合并后文件名
        out_name = _build_merged_filename(group, name_list)
        out_path = output_dir / out_name
        ok = _merge_meal_docx_files(group, out_path)
        if ok:
            merged_paths.append(str(out_path))
            files_to_delete.extend(group)
            # 记录每个旧文件 → 合并后新文件的映射
            for old_file in group:
                merge_map[old_file.name] = str(out_path)
                merge_map[str(old_file)] = str(out_path)
            logger.info("  ✓ 合并完成: %s ← %d 个文件", out_name, len(group))
        else:
            logger.warning("  ✗ 合并失败: %s，保留原文件", out_name)
            for p in group:
                merged_paths.append(str(p))
                merge_map[str(p)] = str(p)
                merge_map[p.name] = str(p)

    # 只备份和删除已合并的生成文件
    for f in files_to_delete:
        try:
            dst = backup_dir / f.name
            if dst.exists():
                stem, suffix = dst.stem, dst.suffix
                i = 1
                while dst.exists():
                    dst = backup_dir / f"{stem}_{i}{suffix}"
                    i += 1
            shutil.copy2(str(f), str(dst))
            f.unlink()
            logger.info("  已备份并删除生成文件: %s", f.name)
        except Exception as exc:
            logger.warning("  备份/删除失败: %s — %s", f.name, exc)

    return {
        "merged_paths": merged_paths,
        "merge_map": merge_map,
        "success": True,
    }


# =========================================================================
# 数据库写入
# =========================================================================

def save_attachment_report(settings, results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """将附件检查结果写入记录数据库。

    预处理（与旧版 save_report 一致）：
    1. 自动补全 附件路径：从 生成文件/匹配附件 推断
    2. 非 PDF/图片 附件自动转 PDF，更新路径（使用多策略转换）
    """
    from invoice_toolkit.database import get_record_db

    if not results:
        return {"success": True, "records_written": 0, "anomalies_written": 0}

    converted_count = 0
    # 跨行去重：同一路径只做一次格式转换，后续直接复用结果
    _fmt_cache: Dict[str, str] = {}

    for row in results:
        # ── 补全附件路径 ──
        att_path = (row.get("附件路径") or "").strip()
        gen_file = (row.get("生成文件") or "").strip()
        match_att = (row.get("匹配附件") or "").strip()

        if not att_path:
            if gen_file:
                att_path = gen_file
            elif match_att and os.path.isabs(match_att):
                att_path = match_att

        # ── 格式转换：确保为 PDF 或图片 ──
        att_path_before = att_path
        if att_path:
            if att_path in _fmt_cache:
                att_path = _fmt_cache[att_path]
            else:
                logger.info("  [格式转换] att_path=%s", att_path[:120])
                att_path_converted = _ensure_upload_format(settings, att_path)
                _fmt_cache[att_path_before] = att_path_converted
                if att_path_converted != att_path_before:
                    converted_count += att_path_converted.count(".pdf")
                att_path = att_path_converted

        if gen_file:
            if gen_file == att_path_before:
                gen_file = att_path
            elif gen_file in _fmt_cache:
                gen_file = _fmt_cache[gen_file]
            else:
                logger.info("  [格式转换] gen_file=%s", gen_file[:120])
                converted = _ensure_upload_format(settings, gen_file)
                _fmt_cache[gen_file] = converted
                gen_file = converted

        row["附件路径"] = att_path
        row["生成文件"] = gen_file

    record_db = get_record_db(settings)

    try:
        record_db.upsert_attachment_check(results)
    except Exception as exc:
        return {"success": False, "error": f"写入数据库失败: {exc}"}

    # 异常项追加到校验详情
    anomaly_count = 0
    for row in results:
        status = row.get("附件状态", "")
        filename = row.get("旧文件名", "")
        category = row.get("附件类别", "")

        if status not in ("缺少附件", "附件校验不通过"):
            continue
        if not filename:
            continue

        if status == "缺少附件":
            missing = row.get("缺少类型", "附件")
            reason = f"[{category}]缺少{missing}"
        else:
            detail = row.get("校验详情", "")
            reason = f"[{category}]附件校验不通过"
            if detail:
                reason = f"{reason}: {detail}"

        try:
            record_db.append_validation_detail(filename, reason)
            anomaly_count += 1
        except Exception:
            pass

    return {
        "success": True,
        "records_written": len(results),
        "anomalies_written": anomaly_count,
        "converted_to_pdf": converted_count,
    }


# =========================================================================
# CLI 入口（供 LLM 通过 bash 调用）
# =========================================================================

def main():
    """命令行入口: python -m scripts.tools <tool_name> <args_json>"""
    import sys
    from invoice_toolkit.config import Settings

    if len(sys.argv) < 2:
        print(json.dumps({"error": "用法: python -m scripts.tools <tool_name> [args_json]"}))
        sys.exit(1)

    tool_name = sys.argv[1]
    args = json.loads(sys.argv[2]) if len(sys.argv) > 2 else {}
    settings = Settings.from_env()

    # 工具路由
    tool_map = {
        "get_config": lambda: get_config(settings),
        "get_ocr_names": lambda: get_ocr_names(settings),
        "collect_files": lambda: collect_files(settings, args["category"]),
        "collect_source_candidates": lambda: collect_source_candidates(settings, args["person"]),
        "lookup_invoice_details": lambda: lookup_invoice_details(settings, args["filename"]),
        "extract_attachment_text": lambda: extract_attachment_text(settings, args["filepath"]),
        "copy_file": lambda: copy_file(args["src"], args["dst_dir"], args.get("mark_used", True)),
        "backup_file": lambda: backup_file(settings, args["filepath"], args.get("delete_original", True)),
        "generate_meal_doc": lambda: generate_meal_doc(
            settings, args["person"], args["amount"], args.get("seller", ""),
            args.get("commodity", ""), args.get("invoice_filename", ""),
            args.get("name_list", settings.NAME_LIST), args.get("reason"),
        ),
        "fix_meal_doc": lambda: fix_meal_doc(
            settings, args["original_path"], args.get("invoice_filename", ""),
            args["person"], args["amount"], args.get("target_persons", []),
            args.get("required_count", 1), args.get("reason_text"),
        ),
        "merge_meal_docs": lambda: merge_meal_docs(settings, args["generated_files"]),
        "save_attachment_report": lambda: save_attachment_report(settings, args["results"]),
    }

    if tool_name not in tool_map:
        print(json.dumps({"error": f"未知工具: {tool_name}"}))
        sys.exit(1)

    try:
        result = tool_map[tool_name]()
        print(json.dumps(result, ensure_ascii=False, default=str))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()